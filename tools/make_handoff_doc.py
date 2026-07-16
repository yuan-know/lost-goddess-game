# -*- coding: utf-8 -*-
"""
生成《失落的女神》项目交接教程 Word 文档。
面向另一台电脑上的 Claude Code —— 读完后它应能:
  1. 完全了解游戏当前进度
  2. 把项目代码下载到那台电脑
  3. 检查并补全 Unity 环境
  4. 在 Unity 中跑起来并能修改代码
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\yuan\lost-goddess-game\失落的女神_项目交接教程_给另一台ClaudeCode.docx"

# ---------- 样式基础 ----------
doc = Document()

# 设定正文默认中文字体
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def set_cn_font(run, name="微软雅黑", size=None, bold=None, color=None):
    run.font.name = name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    set_cn_font(run, "微软雅黑", 18, True, RGBColor(0x1F, 0x3A, 0x5F))
    return p


def h2(text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    set_cn_font(run, "微软雅黑", 14, True, RGBColor(0x2E, 0x5A, 0x88))
    return p


def h3(text):
    p = doc.add_heading(level=3)
    run = p.add_run(text)
    set_cn_font(run, "微软雅黑", 12, True, RGBColor(0x3A, 0x6E, 0xA5))
    return p


def para(text, bold=False, size=10.5, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, "微软雅黑", size, bold, color)
    run.font.italic = italic
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_cn_font(r1, "微软雅黑", 10.5, True)
        r2 = p.add_run(text)
        set_cn_font(r2, "微软雅黑", 10.5, False)
    else:
        r = p.add_run(text)
        set_cn_font(r, "微软雅黑", 10.5, False)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_cn_font(r, "微软雅黑", 10.5, False)
    return p


def code_block(text):
    """灰底等宽代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    # 灰底
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def callout(text, kind="info"):
    """提示框:info=蓝 / warn=橙 / danger=红"""
    colors = {"info": "DCE9F7", "warn": "FCEBD3", "danger": "F8D7DA"}
    fill = colors.get(kind, "DCE9F7")
    label = {"info": "💡 提示", "warn": "⚠️ 注意", "danger": "🚫 重要"}.get(kind, "💡 提示")
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    r1 = p.add_run(label + ": ")
    set_cn_font(r1, "微软雅黑", 10.5, True)
    r2 = p.add_run(text)
    set_cn_font(r2, "微软雅黑", 10.5, False)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htxt)
        set_cn_font(run, "微软雅黑", 10, True, RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            set_cn_font(run, "微软雅黑", 9.5, False)
    if widths:
        for i, w in enumerate(widths):
            for cell in t.columns[i].cells:
                cell.width = Inches(w)
    return t


# ======================================================================
# 封面
# ======================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("《失落的女神》Lost Goddess")
set_cn_font(r, "微软雅黑", 26, True, RGBColor(0x1F, 0x3A, 0x5F))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("项目交接教程 —— 给另一台电脑上的 Claude Code")
set_cn_font(r, "微软雅黑", 14, False, RGBColor(0x55, 0x55, 0x55))

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Unity 2022.3.62f3 · 2D 点击叙事解谜 · 编写日期 2026-07-16")
set_cn_font(r, "微软雅黑", 10, False, RGBColor(0x88, 0x88, 0x88))

doc.add_paragraph()

# 致读者(Claude Code)
callout(
    "你(读这份文档的 Claude Code)运行在【第二台电脑】上,你的搭档是编码员 B。"
    "这份文档的目标:让你读完后能 ① 完全了解游戏当前进度 ② 把项目代码下载到本机 "
    "③ 检查并补全 Unity 环境 ④ 让项目在本机 Unity 里成功运行,并让编码员 B 能修改其中代码。"
    "请按章节顺序逐步执行,每一步都有验证方法。遇到需要真实地址/密码等信息时,主动向编码员 B 索取。",
    "info",
)

doc.add_page_break()

# ======================================================================
# 第 0 章 游戏进度总览
# ======================================================================
h1("第 0 章  游戏进度总览(先读懂项目)")

para("在动手前,先建立对这个项目的完整认知。", size=10.5)

h2("0.1 这是什么游戏")
bullet("类型:类绣湖(Rusty Lake / Cube Escape)风格的 2D 点击式叙事解谜游戏。", )
bullet("独特点:主角老人在场景内【可行走】(区别于经典纯视点绣湖)。行走方式是"
       "「烟火式横版单线」—— 老人只在一条固定水平地平线上左右移动(Y 锁定),点击任意处只取 X。")
bullet("主题:一个探险家追寻失落女神的一生,最终发现寻找的不是神明,而是被遗忘的勇气。")
bullet("主角 = 同一人物的三个年龄形态:青年 Young / 中年 Middle / 老年 Old,对应时间循环 GameState.Era。")
bullet("章节规划:序章 → 老年章 → 中年章 → 青年章(每章 10 天周期,当前做序章)。")

h2("0.2 当前进度(截至 2026-07-16)")
callout("系统层【全套已完成并通过 Play 验证,已 git 提交里程碑】。当前处于"
        "「系统层地基稳固、准备进入内容层」的阶段。", "info")
bullet("接口契约 §1~§10 全部实现(状态全标 🟢),21 个 C# 脚本,batchmode 编译零错误。")
bullet("三形态立绘(青/中/老)已从美术分层 PSD 导出、导入 Unity 并 Play 验证通过"
       "(走路、大小、边缘白线问题全部解决)。")
bullet("验证沙盒 Sandbox.unity 已生成,可 Play 跑通全链路:点空地走 / 捡提灯入包 / "
       "持灯开门切房间 / F5 存档 F9 读档 / 按 1·2·3 实时切换三形态立绘。")
bullet("git 里程碑已提交:commit e69af0b(系统层+立绘+沙盒)、4f4c638(接口契约状态更新)。")

h3("已交付的 21 个脚本(系统层)")
table(
    ["目录", "脚本", "职责"],
    [
        ["Save/", "SaveData / GameState / SaveSystem", "存档数据结构、全局状态静态 API、JSON 读写到 persistentDataPath"],
        ["Managers/", "GameManager", "单例 + 协程宿主 + 按 Era 进章"],
        ["Managers/", "SceneLoader", "房间切换协程(淡入淡出)+ 程序化房间构建双路加载"],
        ["Managers/", "AudioManager", "从 Resources/Audio 加载,缺文件静默"],
        ["Managers/", "InventorySystem", "背包:选中物品→点目标物,物品存 GameState"],
        ["Managers/", "ClickInputManager", "旧版 Input + Physics2D.OverlapPoint 点击调度"],
        ["Player/", "PlayerController", "横版点击寻路(只左右)、朝向翻转、Animator 参数接口、脚底 Y 深度排序、Teleport"],
        ["Player/", "WalkableArea", "烟火式横版可走线:地平线 Y + minX/maxX(青色 Gizmo)"],
        ["Interaction/", "InteractableBase", "可交互物基类:走近 interactPoint → OnClick,含 eraRestricted 能力钩子/高亮/持物判定"],
        ["Interaction/", "PuzzleBase", "谜题基类:Solve/OnSolved/存档恢复"],
        ["Interaction/", "CloseupView", "运行时特写层 UI + 锁定行走"],
        ["Dialogue/", "DialogueSystem", "按 id 播对白,运行时底部对白条 UI,点击推进"],
        ["Utils/", "Keys", "Flags/Items/Rooms/Dialogues/Sfx/Bgm 常量集中管理"],
        ["Utils/", "SandboxBootstrap", "程序化占位房间 + 按 Era 加载立绘 + 按 1/2/3 切形态"],
        ["Utils/Editor/", "SandboxSceneCreator", "菜单生成验证沙盒场景"],
        ["Utils/Editor/", "CharacterSpriteImporter", "AssetPostprocessor:自动配置立绘 PPU/锚点/mipmap"],
        ["Utils/Editor/", "ForceReimportCharacters", "菜单强制重导入立绘"],
        ["Rooms/Prologue/", "Interact_Door / Interact_PickupLamp", "示例交互物(门 / 捡提灯)"],
    ],
    widths=[1.1, 2.2, 3.7],
)

h2("0.3 关键技术决策(改代码前必须知道)")
bullet("输入:用【旧版 Input + Physics2D】,不用 InputSystem(虽然装了包)。", "输入方案 — ")
bullet("淡入淡出:用协程实现,【不引 DOTween】(免费 UPM 源没有,行走原型不需要)。", "转场 — ")
bullet("形态切换:【分章固定】—— 进章时按 Era 选 prefab,不做运行时热切换(沙盒里的 1/2/3 只是调试对比用)。", "三形态 — ")
bullet("专属能力(青年蛮力/中年组合/老年提灯):先不做,只在 InteractableBase 留 eraRestricted 钩子。", "能力 — ")
bullet("立绘导入:PPU=600(人物占屏约 1/3)、锚点 BottomCenter(脚底居中)、mipmap 开 + 三线性过滤。", "美术 — ")

h2("0.4 必读文档(项目内 docs/ 目录)")
bullet("docs/接口契约.md —— 系统层 A 提供给内容层 B 的接口清单,§1~§10。【改代码前必读】。", "★ ")
bullet("docs/程序开发计划方案.md —— 双人并行开发方案、目录结构、协作铁律、排期。", )
bullet("docs/人物设定_三形态.md —— 主角青/中/老三形态的美术与能力设定。", )
bullet("README.md —— 环境要求与团队铁律速查。", )

doc.add_page_break()

# ======================================================================
# 第 1 章 下载项目代码
# ======================================================================
h1("第 1 章  把项目代码下载到本机")

callout("分发方式:通过远程 git 仓库(Gitee 或 GitHub)。项目使用 Git LFS 管理美术图片"
        "(目前 3 张三形态立绘),所以【必须先装并初始化 LFS,否则拉下来的图片是文本指针而非真实图片】。", "warn")

h2("1.1 前置:向编码员 B 索取仓库地址")
para("这份文档写作时,A 方(第一台电脑)可能尚未把项目推到远程。请先确认:")
numbered("向编码员 B / A 索取远程仓库地址(形如 https://gitee.com/xxx/lost-goddess-game.git)。")
numbered("如果 A 方还没推送,请提醒 A 方先在第一台电脑执行推送(推送命令见本章末「附:A 方推送备忘」)。")
numbered("确认你是否有该仓库的访问权限(私有仓库需要账号登录或 access token)。")

h2("1.2 检查 git 与 git-lfs")
code_block("git --version\ngit lfs version")
para("如果 git lfs 未安装:")
bullet("Windows:从 https://git-lfs.com 下载安装,或 winget install GitHub.GitLFS")
bullet("装好后执行一次:git lfs install(为当前用户启用 LFS 钩子)")

h2("1.3 克隆项目")
para("在你想放项目的目录下执行(把 <仓库地址> 换成 B 给你的真实地址):")
code_block(
    "git lfs install\n"
    "git clone <仓库地址> lost-goddess-game\n"
    "cd lost-goddess-game\n"
    "git lfs pull        # 确保 LFS 图片被真正拉下来"
)

h2("1.4 验证下载完整")
para("克隆后逐项确认:")
code_block(
    "# 1) 提交历史应看到里程碑提交\n"
    "git log --oneline -5\n\n"
    "# 2) LFS 图片应是真实文件(每张 1MB 上下),不是几百字节的文本指针\n"
    "git lfs ls-files\n"
    "#    应列出 3 张:young.png / middle.png / old.png\n\n"
    "# 3) 三张立绘应存在且体积正常\n"
    "ls -lh Assets/_Project/Resources/Characters/"
)
callout("如果 young/middle/old.png 只有几百字节,说明 LFS 没生效 —— 重新执行 "
        "git lfs install 再 git lfs pull。", "warn")

h2("1.5 目录结构速览(下载后应看到)")
code_block(
    "lost-goddess-game/\n"
    "├── Assets/_Project/\n"
    "│   ├── Scenes/Sandbox.unity          ← 验证沙盒(第 3 章要打开它)\n"
    "│   ├── Scripts/                       ← 21 个系统层脚本\n"
    "│   │   ├── Core/ (Save/Managers/Interaction/Player/Dialogue)\n"
    "│   │   ├── Rooms/Prologue/\n"
    "│   │   └── Utils/ (含 Editor 子目录)\n"
    "│   └── Resources/Characters/          ← 三形态立绘(LFS)\n"
    "├── Packages/manifest.json             ← 包依赖清单(第 2 章要核对)\n"
    "├── ProjectSettings/ProjectVersion.txt ← Unity 版本锁定\n"
    "├── docs/                              ← 4 份必读文档\n"
    "└── tools/                             ← 立绘导出 Python 管线脚本"
)

h2("附:A 方(第一台电脑)推送备忘")
para("(此节给 A 方参考。若项目已在远程,可跳过。)", italic=True)
code_block(
    "# 在第一台电脑的项目根目录执行\n"
    "git remote add origin <仓库地址>\n"
    "git push -u origin main\n"
    "git lfs push origin main --all     # 确保 LFS 大文件也推上去"
)

doc.add_page_break()

# ======================================================================
# 第 2 章 检查并补全 Unity 环境
# ======================================================================
h1("第 2 章  检查并补全 Unity 环境(重点)")

callout("编码员 B 已装好与 A 相同版本的 Unity Hub 与 Unity 2022.3,但【环境可能不完整】。"
        "本章请你逐项检查并补全,确保能编译运行。", "danger")

h2("2.1 Unity 版本必须完全一致")
para("项目锁定版本(见 ProjectSettings/ProjectVersion.txt):")
code_block("m_EditorVersion: 2022.3.62f3\nm_EditorVersionWithRevision: 2022.3.62f3 (96770f904ca7)")
bullet("必须用 2022.3.62f3 打开,【不要】用其他子版本(如 2022.3.10)或 Unity 6,否则会触发资源升级、产生难以合并的差异。")
bullet("检查方法:打开 Unity Hub → Installs,确认列表里有 2022.3.62f3。")
bullet("若没有:Unity Hub → Installs → Install Editor → Archive,精确安装 2022.3.62f3。")
callout("这台机器上另装有 Unity 6000.0.17f1 的话,别用错;本项目只认 2022.3.62f3。", "warn")

h3("必装的 Unity 模块(安装 Editor 时勾选)")
bullet("Windows Build Support (IL2CPP / Mono) —— 打包用")
bullet("2D 相关模块(通常随 2022.3 内置,但确认已启用)")

h2("2.2 包依赖(Package)—— 核对 manifest.json")
para("项目的 Packages/manifest.json 已锁定以下依赖。用 Unity 打开项目时,Package Manager 会"
     "根据这个文件【自动下载还原】这些包 —— 这就是为什么「环境不完整」通常能被自动修复。"
     "但需联网,且首次还原耗时。请核对下表:")
table(
    ["包名", "版本", "用途"],
    [
        ["com.unity.2d.animation", "9.1.3", "2D 骨骼动画(后续接骨骼动画用)"],
        ["com.unity.2d.sprite", "1.0.0", "Sprite 编辑"],
        ["com.unity.2d.tilemap", "1.0.0", "瓦片地图"],
        ["com.unity.inputsystem", "1.7.0", "新输入系统(已装但当前用旧版 Input)"],
        ["com.unity.nuget.newtonsoft-json", "3.2.1", "JSON 序列化(存档用)"],
        ["com.unity.textmeshpro", "3.0.9", "文字渲染(对白/UI)"],
        ["com.unity.ugui", "1.0.0", "UI 系统"],
        ["modules.audio / animation / imgui /", "1.0.0", "Unity 内置模块"],
        ["  jsonserialize / ui / uielements /", "1.0.0", "(随引擎)"],
        ["  physics2d / unitywebrequest", "1.0.0", "physics2d 是点击命中的关键"],
    ],
    widths=[2.6, 1.0, 3.4],
)
para("检查与补全步骤:")
numbered("用 Unity 2022.3.62f3 打开项目(Unity Hub → Open → 选项目根目录)。")
numbered("首次打开会自动 Import 并还原 Packages,耐心等待(状态栏有进度)。")
numbered("打开 Window → Package Manager → 左上下拉选「In Project」,核对上表包都在且版本一致。")
numbered("若某包缺失/报错:Package Manager 里手动 Add(按名+版本),或直接确认 manifest.json 内容与上表一致后重启 Unity 触发还原。")

callout("Newtonsoft JSON(存档核心依赖)如果没还原成功,会导致 SaveSystem 编译报错。"
        "这是最常见的「环境不完整」症状 —— 确认 com.unity.nuget.newtonsoft-json 已在 In Project 列表。", "warn")

h2("2.3 Unity 工程设置(应随仓库带入,确认即可)")
para("以下设置已在 ProjectSettings 里入库,理论上克隆后自动生效,但请确认:")
bullet("Edit → Project Settings → Editor → Asset Serialization = Force Text(场景文本化,防合并冲突)")
bullet("Edit → Project Settings → Editor → Version Control Mode = Visible Meta Files")

h2("2.4 立绘导入设置(可能需要手动触发一次 Reimport)")
para("三张立绘依赖 CharacterSpriteImporter.cs(AssetPostprocessor)自动配置 PPU=600、锚点脚底居中、"
     "mipmap 开等参数。克隆后 .meta 通常已带正确参数;若发现立绘显示异常(过大/过小/边缘白线/发糊):")
numbered("Project 窗口定位 Assets/_Project/Resources/Characters 文件夹。")
numbered("右键 → Reimport,让 CharacterSpriteImporter 重新套用参数。")
numbered("或用菜单栏「失落的女神 ▸ 强制重导入立绘」(ForceReimportCharacters 提供)。")

h2("2.5 命令行验证编译(可选,推荐你执行)")
para("你(Claude Code)可以用 batchmode 无界面编译,快速确认脚本零错误。把路径按本机 Unity 安装位置调整:")
code_block(
    '"<Unity安装路径>/Editor/Unity.exe" -batchmode -quit -nographics \\\n'
    '  -projectPath "<项目根目录>" \\\n'
    '  -logFile "C:/temp/compile.log" -accept-apiupdate\n\n'
    "# 退出码 0 = 编译成功。例(A 方的路径,供参考):\n"
    '# "D:/unity/2022.3.62f3/Editor/Unity.exe" -batchmode -quit -nographics \\\n'
    '#   -projectPath "C:/.../lost-goddess-game" -logFile "compile.log" -accept-apiupdate'
)
callout("batchmode 编译时【必须先关闭 Unity 编辑器】,否则工程被锁,退出码非 0(常见 EXIT=21)。", "warn")

doc.add_page_break()

# ======================================================================
# 第 3 章 运行验证
# ======================================================================
h1("第 3 章  在 Unity 中运行验证")

para("环境就绪后,用验证沙盒确认项目能跑起来。")

h2("3.1 打开验证沙盒并 Play")
numbered("Unity 里打开场景:Assets/_Project/Scenes/Sandbox.unity(双击)。")
numbered("如果沙盒场景不存在或想重建:菜单栏「失落的女神 ▸ 生成验证沙盒场景」。")
numbered("点顶部 ▶ Play 按钮运行。")

h2("3.2 逐项验证功能(全部通过 = 项目在本机跑通)")
table(
    ["操作", "预期结果"],
    [
        ["点击空地", "老人(默认老年立绘)沿地平线左右走过去,朝向随方向翻转"],
        ["点击提灯(黄色物件)", "老人走近 → 捡起提灯,进入背包"],
        ["点击门(棕色物件)", "老人走近 → 若持灯则开门并切换房间"],
        ["按 F5 / F9", "F5 存档 / F9 读档(存到 persistentDataPath)"],
        ["按 1 / 2 / 3", "实时切换青年 / 中年 / 老年三形态立绘,对比走路与大小"],
    ],
    widths=[2.6, 4.4],
)
callout("建议把 Game 窗口分辨率设为 Full HD(1920×1080),不要用 Free Aspect + 低缩放,"
        "否则立绘看起来发糊(这是显示分辨率问题,不是资源问题)。", "info")

h2("3.3 若运行报错")
bullet("控制台报 Newtonsoft / Json 相关错误 → 回第 2.2 节确认包还原。")
bullet("立绘不显示或异常 → 回第 2.4 节 Reimport。")
bullet("点击无反应 → 确认 physics2d 模块在,且 Game 窗口有焦点。")
bullet("脚本编译错误 → 看 Console 具体行,多半是包缺失或 Unity 版本不符。")

doc.add_page_break()

# ======================================================================
# 第 4 章 如何修改代码
# ======================================================================
h1("第 4 章  如何修改代码(交给编码员 B)")

h2("4.1 分工:A 系统层 / B 内容层")
table(
    ["角色", "代号", "负责", "地盘(目录)"],
    [
        ["编码员 A", "系统程序", "框架/存档/场景切换/交互谜题基类/各 Manager", "Scripts/Core/"],
        ["编码员 B", "内容程序", "具体房间/可交互物件/谜题逻辑/剧情触发/过场", "Scripts/Rooms/ 、Scripts/Puzzles/"],
    ],
    widths=[1.2, 1.2, 3.2, 1.4],
)

h2("4.2 B 做一个房间的标准姿势")
numbered("在 Scripts/Rooms/<章节>/ 下新建交互脚本,继承 InteractableBase,实现 OnClick()。")
numbered("房间场景里摆一个 WalkableArea,设好地平线 Y 与左右边界 minX/maxX(Scene 视图有青色地平线 Gizmo)。")
numbered("每个可交互物摆一个 interactPoint 空物体标「老人站哪儿点它」,其 Y 对齐地平线。")
numbered("用 GameState 存开关/物品,用 SceneLoader.GoToRoom() 切场景,用 DialogueSystem.Show() 播对白。")
numbered("Flag / 物品 key 一律写进 Utils/Keys.cs,禁止散落魔法字符串。")

h3("交互流程要点(本作主角可行走)")
para("玩家点可交互物 → 系统自动让老人走到该物 interactPoint → 到位后才回调 OnClick()。"
     "B 只写 OnClick() 里「到位后发生什么」。远处不需走近的物件设 walkToBeforeInteract=false 立即触发。")

h2("4.3 协作铁律(防止 Unity 冲突)")
bullet("一个场景 / 预制体同一时间只有一个人改,开工前在群里/看板认领。", "① ")
bullet("系统层公共文件(Core/ 下)只有 A 改;B 需要新全局功能 → 提给 A,别自己动。", "② ")
bullet("每天开工先 git pull,收工前 git push,别攒一周再合。", "③ ")
bullet("提交信息带前缀:[core] / [room] / [puzzle] / [art] / [fix]。", "④ ")
bullet("接口若要改,A 先改 docs/接口契约.md 并通知 B。B 只调用契约里的公开接口,不读写系统层内部字段。", "⑤ ")

h2("4.4 日常 git 流程(B 在第二台电脑)")
code_block(
    "git pull                          # 每天开工第一件事\n"
    "git checkout -b feat/room-hall    # 一个房间一条分支\n"
    "# ... 改代码 ...\n"
    'git add . && git commit -m "[room] hall 谜题完成"\n'
    "git push -u origin feat/room-hall # 或按团队约定推 dev"
)
callout("提交含图片时,确认 git lfs 正常工作(git lfs status 能看到图片走 LFS)。", "info")

h2("4.5 下一步开发候选(A 方规划,供 B 参考)")
bullet("序章第一个真实房间(需策划提供房间清单 + 谜题清单)。")
bullet("三形态 prefab(青年/中年空壳),为分章切换铺路。")
bullet("骨骼动画:加 com.unity.2d.psdimporter 包 → PSD 原样导入 → 美术在 Skinning Editor 绑骨骼、"
       "摆 idle/walk → 建 Animator Controller 接 PlayerController.animator(绑骨骼是美术手工活,程序备环境)。")
bullet("背包 / 特写 UI 打磨。")

doc.add_page_break()

# ======================================================================
# 第 5 章 快速检查清单
# ======================================================================
h1("第 5 章  快速检查清单(照此逐项打勾)")

h2("下载阶段")
bullet("☐ 已向 B 拿到远程仓库地址并确认访问权限")
bullet("☐ git lfs install 已执行")
bullet("☐ git clone 成功,git log 能看到里程碑提交")
bullet("☐ git lfs ls-files 列出 3 张立绘,且文件体积正常(约 1MB)")

h2("环境阶段")
bullet("☐ Unity Hub 里有 2022.3.62f3(与 ProjectVersion.txt 一致)")
bullet("☐ 用 2022.3.62f3 打开项目,Packages 自动还原完成")
bullet("☐ Package Manager「In Project」核对依赖齐全(尤其 Newtonsoft JSON)")
bullet("☐ 无脚本编译错误(Console 无红字 / batchmode 退出码 0)")

h2("运行阶段")
bullet("☐ 打开 Sandbox.unity 并 Play")
bullet("☐ 点空地老人会走 / 捡提灯 / 持灯开门 / F5F9 存读档 / 1·2·3 切三形态 全部正常")

h2("协作阶段")
bullet("☐ 已读 docs/接口契约.md 和 docs/程序开发计划方案.md")
bullet("☐ 理解 A/B 分工与 5 条协作铁律")
bullet("☐ B 能新建分支、提交、推送")

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("—— 全部打勾 = 交接完成,项目已可在第二台电脑运行并修改 ——")
set_cn_font(r, "微软雅黑", 11, True, RGBColor(0x1F, 0x3A, 0x5F))

doc.save(OUT)
print("已生成:", OUT)
