# -*- coding: utf-8 -*-
"""把 docs/骨骼动画操作指南.md 的内容生成一份排版精美的 Word。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\yuan\lost-goddess-game\docs\骨骼动画操作指南.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def cn(run, name="微软雅黑", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color


def h1(t):
    p = doc.add_heading(level=1); r = p.add_run(t)
    cn(r, "微软雅黑", 17, True, RGBColor(0x1F, 0x3A, 0x5F)); return p

def h2(t):
    p = doc.add_heading(level=2); r = p.add_run(t)
    cn(r, "微软雅黑", 13.5, True, RGBColor(0x2E, 0x5A, 0x88)); return p

def h3(t):
    p = doc.add_heading(level=3); r = p.add_run(t)
    cn(r, "微软雅黑", 11.5, True, RGBColor(0x3A, 0x6E, 0xA5)); return p

def para(t, bold=False, size=10.5, color=None, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t)
    cn(r, "微软雅黑", size, bold, color); r.font.italic = italic; return p

def bullet(t, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if prefix:
        r1 = p.add_run(prefix); cn(r1, "微软雅黑", 10.5, True)
        r2 = p.add_run(t); cn(r2, "微软雅黑", 10.5, False)
    else:
        r = p.add_run(t); cn(r, "微软雅黑", 10.5, False)
    return p

def num(t):
    p = doc.add_paragraph(style="List Number"); r = p.add_run(t)
    cn(r, "微软雅黑", 10.5, False); return p

def code(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    pPr = p._element.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd); return p

def callout(t, kind="info"):
    colors = {"info": "DCE9F7", "warn": "FCEBD3", "danger": "F8D7DA", "gold": "FFF4CE"}
    label = {"info": "💡 提示", "warn": "⚠️ 注意", "danger": "🚫 重要", "gold": "★ 黄金法则"}
    p = doc.add_paragraph(); pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), colors.get(kind, "DCE9F7")); pPr.append(shd)
    r1 = p.add_run(label.get(kind, "💡 提示") + ": "); cn(r1, "微软雅黑", 10.5, True)
    r2 = p.add_run(t); cn(r2, "微软雅黑", 10.5, False); return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""; r = hdr[i].paragraphs[0].add_run(h)
        cn(r, "微软雅黑", 10, True, RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""; r = cells[i].paragraphs[0].add_run(str(v))
            cn(r, "微软雅黑", 9.5, False)
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells: c.width = Inches(w)
    return t


# 封面
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("骨骼动画操作指南"); cn(r, "微软雅黑", 24, True, RGBColor(0x1F, 0x3A, 0x5F))
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("《失落的女神》老年形态 idle / walk"); cn(r, "微软雅黑", 14, False, RGBColor(0x55, 0x55, 0x55))
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Unity 2022.3.62f3 · 2D Animation 骨骼 · 2026-07-16")
cn(r, "微软雅黑", 10, False, RGBColor(0x88, 0x88, 0x88))
doc.add_paragraph()

callout("本指南覆盖必须在 Unity 编辑器里手工完成的部分(绑骨骼、蒙皮、摆关键帧、接线)。"
        "能自动化的部分(装包、PSD 分层导入、Animator 状态机)已由程序做成菜单一键工具,"
        "指南会告诉你在哪一步点它们。目标:让老年老人从静止立绘变成待机微动 / 行走迈腿摆臂的"
        "骨骼动画,并接到 PlayerController,点击行走时自动播 walk、停下播 idle。", "info")

para("术语:Bone=骨骼,Skinning/Weight=蒙皮/权重(把图层像素绑到骨头),Rig=骨架,Clip=动画片段。",
     italic=True, size=9.5, color=RGBColor(0x77, 0x77, 0x77))

# 阶段0
h1("阶段 0  环境就绪(程序已做,你只需确认)")
num("包已装:manifest.json 已加 com.unity.2d.psdimporter,切回 Unity 自动还原。")
num("确认:Window ▸ Package Manager → 左上下拉 In Project → 能看到 2D PSD Importer 和 2D Animation。")
num("PSD 已在项目:Assets/_Project/Art/Characters/characters_3forms.psd(青/中/老三组,每组12部位层+1完整立绘层)。")
callout("如果 Package Manager 里没有 2D PSD Importer:包没还原成功。确认 manifest.json 有 "
        '"com.unity.2d.psdimporter": "8.1.1",然后 Assets ▸ Reimport All 或重启 Unity。', "warn")

# 阶段1
h1("阶段 1  把 PSD 配成 Mosaic 骨骼导入(程序菜单,一键)")
para("点菜单:", bold=True)
callout("失落的女神 ▸ 配置角色PSD(Mosaic骨骼导入)", "info")
para("它会把 PSD 设成:Mosaic(各图层拆成独立 Sprite)+ Import Hidden(导入隐藏层,"
     "本 PSD 部位层默认隐藏,必须开否则导进来是空的)+ Character 模式(可绑骨骼)+ "
     "PPU 600 + 锚点脚底,然后自动 Reimport。")
para("成功后,在 Project 窗口点开 characters_3forms.psd 左边的小三角,应能看到子 Sprite:"
     "老年头脖子 / 老年躯干大衣 / 老年右大臂肘关节…… 每组 12 个,共 3 组。")
callout("只做老年:三组都会被导入,骨骼动画本次只处理老年 12 个部位,青/中年 Sprite 先不管。", "info")

# 阶段2
h1("阶段 2  进入 Skinning Editor 绑骨骼(手工核心)")
num("Project 里选中 characters_3forms.psd(选整个 psd,不是子 Sprite)。")
num("Inspector 点 Sprite Editor 按钮 → 打开 Sprite Editor 窗口。")
num("窗口左上角下拉(默认 Sprite Editor)→ 切到 Skinning Editor。画布上应显示拼成完整人形的老人。")

h2("2.1 造骨骼(Create Bone)")
num("左侧工具栏点 Create Bone(骨头图标)。")
num("从盆骨/腰部开始,一根根画(点起点→点终点定一根;继续点接着往下延伸成父子链)。")
para("老年人形推荐骨架(拄拐、佝偻,可简化):")
code(
    "Root(盆骨,两腿之间靠上)\n"
    "├─ Spine 脊柱(向上到胸)\n"
    "│   ├─ Head 头(脖子→头)\n"
    "│   ├─ Arm_L_Upper 左大臂 → Arm_L_Lower 左小臂手\n"
    "│   └─ Arm_R_Upper 右大臂 → Arm_R_Lower 右小臂手\n"
    "├─ Leg_L_Upper 左大腿 → Leg_L_Lower 左小腿 → Foot_L 左脚\n"
    "└─ Leg_R_Upper 右大腿 → Leg_R_Lower 右小腿 → Foot_R 右脚"
)
bullet("画骨头要接着关节位置画(肩、肘、髋、膝、踝),后面弯曲才自然。")
bullet("另起一条不连续骨链:先按 Esc 断开,再点新起点。")
bullet("画错用 Edit Bone 拖动调整,或选中删除重画。")

h2("2.2 蒙皮(把图层像素绑到骨头)")
num("点 Auto Geometry → Generate For All Visible,让每个部位生成网格(参数用默认,Subdivide 给 0~2)。")
num("点 Auto Weights → Generate For All Visible,自动把网格权重分给最近骨头。")
num("检查权重:Pose 模式拖动骨头,看部位是否跟对的骨头动(手臂动躯干不该变形)。")
bullet("有问题用 Weight Brush:选中目标骨头,在出问题部位上涂,增/减该骨头影响。")

h2("2.3 摆静止基准姿势(推荐)")
bullet("Preview Pose 里把老人摆成自然拄拐站立(略佝偻),作为 idle 基准。")
callout("摆好后 Skinning Editor 顶部点 Apply 保存 rig。关窗口前必须 Apply,否则骨骼/权重不保存!", "danger")

# 阶段3
h1("阶段 3  把角色拖进场景,准备做动画")
num("characters_3forms.psd 从 Project 拖进 Hierarchy —— 因开了 Generate GO Hierarchy,"
    "会生成带完整骨骼层级的 GameObject(Sprite Renderer + Sprite Skin + 骨骼 Transform)。")
num("对象里三组都在;把青年、中年子物体 SetActive(false) 或删掉,只留老年,重命名 OldMan_Rigged。")
num("拖回 Assets/_Project/Prefabs/ 做成 OldMan.prefab。")
callout("拖进去发现部位错位/散开:多半 Auto Geometry/Weights 没做全,回阶段 2 补。", "warn")

# 阶段4
h1("阶段 4  生成 Animator 状态机(程序菜单,一键)")
para("点菜单:", bold=True)
callout("失落的女神 ▸ 生成老年Animator脚手架", "info")
para("生成 Assets/_Project/Art/Animations/Old/OldMan.controller,内含:")
bullet("参数 isWalking(bool)—— 与 PlayerController.walkParam 一致")
bullet("状态 Idle(默认)/ Walk,已连好双向过渡")
bullet("占位 clip Old_Idle(不循环)/ Old_Walk(循环)")
para("之后你只需给这两个 clip 摆帧(阶段 5),不用自己连状态机。")

# 阶段5
h1("阶段 5  摆 idle / walk 关键帧(手工核心)")
num("选中场景里的 OldMan_Rigged(或双击 Prefab 编辑)。")
num("给它加 Animator 组件(若没自带),Controller 拖成 OldMan.controller。")
num("打开 Window ▸ Animation ▸ Animation 窗口(不是 Animator!Animation 才是摆帧的)。左上 clip 下拉选 Old_Idle。")

h3("5.1 摆 Idle(待机微动 + 手抖)")
bullet("Idle 建议 1~2 秒循环。老年设定:待机时手微微颤抖,身体轻微起伏。")
bullet("点录制键(红点),0 帧:选中骨骼,记录当前角度打关键帧。")
bullet("中间帧(0.5s):手/小臂骨头旋转极小角度(1~3°),身体略下沉。")
bullet("末帧回到与 0 帧相同姿势(循环平滑),关掉录制。Old_Idle 已设循环。")

h3("5.2 摆 Walk(拖拽步态)")
bullet("切 clip 下拉到 Old_Walk。老年:擦地拖拽的缓慢步态,幅度小。")
para("最简 4 帧循环(0 / 0.25 / 0.5 / 0.75s,总长约 1s):")
bullet("0s:左腿在前(略抬)、右腿在后;对侧手臂前后摆(幅度小,老年迟缓)。")
bullet("0.25s:双腿接近并拢过渡。")
bullet("0.5s:右腿在前、左腿在后(与 0s 镜像)。")
bullet("0.75s:并拢过渡。(循环回 0s)")
bullet("拄拐:拐杖随左手小幅前后点地。身体整体略上下起伏(比青年小)。")
callout("摆帧技巧:先只摆大腿/小腿两三根骨头把迈步做出来,跑起来看感觉,再加手臂和身体起伏。"
        "别一次调 12 根骨头。", "info")

# 阶段6
h1("阶段 6  接线到 PlayerController(手工,几下点击)")
para("要让点击行走 → 自动切 walk/idle,需把 Animator 接到 PlayerController。")
h3("方式 A:直接用 OldMan.prefab 当玩家(推荐先验证)")
num("给 OldMan_Rigged 加 PlayerController 组件(若没有)。")
para("Inspector 里 PlayerController 设置:")
bullet("Animator 字段 → 拖入本对象的 Animator 组件。")
bullet("Walk Param → 保持 isWalking(默认,别改)。")
bullet("Sprite Faces Right → 老年立绘朝左,设 false(与沙盒一致)。")
bullet("Sorting Target → 拖入躯干 SpriteRenderer(或给根加 SortingGroup 后拖它)。")
bullet("moveSpeed 老年建议 2.2。")
num("Play:点空地老人走 → 自动播 Old_Walk;到达停下 → 回 Old_Idle。")

h3("方式 B:接到现有沙盒")
bullet("沙盒当前用 Resources/Characters/old.png 整张立绘(无骨骼)。")
bullet("要让沙盒改用骨骼 Prefab,需程序改 SandboxBootstrap.BuildPlayer 改成实例化 OldMan.prefab。"
       "这一步交给程序(A);你先用方式 A 在空场景验证骨骼动画跑通即可。")

# 验收
h1("验收标准(全达成 = 老年骨骼动画完成)")
for t in [
    "PSD 在 Project 里能展开看到老年 12 个部位 Sprite",
    "Skinning Editor 里骨骼绑好、Pose 拖动骨头人物正确跟随变形",
    "Old_Idle 播放时老人待机微动/手抖,循环平滑",
    "Old_Walk 播放时老人迈腿摆臂拖拽步态,循环平滑",
    "挂 PlayerController 后,Play 点击行走自动 walk、停下自动 idle",
    "朝向翻转正常,脚底不穿地平线",
]:
    bullet("☐ " + t)

# 常见问题
h1("常见问题")
table(
    ["现象", "原因 / 解决"],
    [
        ["PSD 展开没有子 Sprite", "没配 Mosaic / Import Hidden。点菜单「配置角色PSD」。"],
        ["导进来部位是空的/缺件", "Import Hidden 没开(本 PSD 部位层默认隐藏)。同上重配。"],
        ["拖进场景部位散开错位", "Auto Geometry/Weights 没做全,回阶段 2.2。"],
        ["拖骨头某部位不跟随", "该部位权重没分到骨头,用 Weight Brush 补,或重跑 Auto Weights。"],
        ["关窗口后骨骼没了", "没点 Apply。Skinning Editor 改动必须 Apply 才保存。"],
        ["Play 不切 walk 动画", "animator 没拖 / walkParam 不是 isWalking / Controller 没挂。"],
        ["动画播了但人物不移动", "位移由 PlayerController 代码驱动,别在 walk clip 里 K 根节点位移。"],
    ],
    widths=[2.3, 4.5],
)
callout("位移(走多远)= PlayerController 代码负责;姿态(腿怎么迈)= 动画负责。"
        "walk clip 里只摆骨头相对旋转/摆动,不要移动角色根节点位置。", "gold")

doc.save(OUT)
print("已生成:", OUT)
