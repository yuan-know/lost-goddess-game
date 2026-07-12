# -*- coding: utf-8 -*-
"""把《失落的女神》程序开发计划方案转成简易可读的 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- 全局中文字体 ----------
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

ACCENT = RGBColor(0x6B, 0x4E, 0x8E)   # 神庙紫
DARK = RGBColor(0x33, 0x33, 0x33)

def set_font(run, size=10.5, bold=False, color=None, name='Microsoft YaHei'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    if level == 1:
        run = p.add_run(text)
        set_font(run, 16, True, ACCENT)
        p.space_after = Pt(6)
        # 底部横线
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '6B4E8E')
        pbdr.append(bottom); pPr.append(pbdr)
    else:
        run = p.add_run(text)
        set_font(run, 13, True, DARK)
    return p

def add_body(text, bullet=False, indent=0):
    p = doc.add_paragraph(style='List Bullet' if bullet else None)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, 10.5)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 4'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].clear()
        run = hdr[i].paragraphs[0].add_run(h)
        set_font(run, 10, True, RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].clear()
            run = cells[i].paragraphs[0].add_run(val)
            set_font(run, 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ================= 封面 =================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《失落的女神》')
set_font(run, 26, True, ACCENT)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('程序开发计划方案 · 双人并行')
set_font(run, 14, False, DARK)

quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = quote.add_run('“一个老人追寻失落女神的旅程,\n最终发现自己寻找的不是神明,而是被遗忘的勇气。”')
set_font(run, 11, False, RGBColor(0x88, 0x88, 0x88))
run.italic = True

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('类绣湖 2D 点击式叙事解谜(主角可行走)  |  Unity 2022.3.62f3  |  Git + Git LFS  |  程序 ×2')
set_font(run, 10, False, DARK)
doc.add_paragraph()

# ================= 一、核心理念 =================
add_heading('一、核心理念:为什么绣湖类适合双人并行', 1)
add_body('绣湖类游戏本质 = 一堆相对独立的房间/场景 + 若干可交互物件 + 谜题状态机 + 全局存档/循环。这种结构天然适合把代码拆成两层:')
add_body('系统层(横向、全局、一次做好长期复用)', bullet=True)
add_body('内容层(纵向、每个房间/谜题一份、可无限并行)', bullet=True)
add_body('只要系统层的接口先约定好,两位程序员就能各做各的房间/谜题,几乎不踩对方代码。这就是本方案分工的核心逻辑。')

# ================= 二、分工模式 =================
add_heading('二、分工模式(推荐:系统层 A + 内容层 B)', 1)
add_table(
    ['角色', '代号', '负责', '一句话职责'],
    [['程序员 A', '系统程序', '底层框架、存档、场景切换、时间循环、交互/谜题基类、各 Manager', '搭舞台、定规则'],
     ['程序员 B', '内容程序', '具体房间、可交互物件、谜题逻辑、剧情触发、过场', '在舞台上演戏']]
)
add_heading('为什么这样分而不是"一人一半房间"', 2)
add_body('序章阶段系统层还没成型,若两人都做房间,会同时抢着改 GameManager、存档等公共文件 → 天天冲突。', bullet=True)
add_body('先由 A 把地基和"积木接口"打好,B 立刻能用这些积木拼房间,互不干扰。', bullet=True)
add_body('中后期自然过渡到并行:系统层稳定后,A 也可下场做房间,两人各包不同房间,纯内容并行,零冲突。', bullet=True)
add_heading('阶段性并行度', 2)
add_body('序章前 3 天: A 打地基     |  B 做美术占位/工具/第一个交互原型(弱依赖)', indent=0.2)
add_body('序章中 4 天: A 补系统     |  B 大量产出房间与谜题(强并行)', indent=0.2)
add_body('序章后 3 天: A、B 各包剩余房间 + 联调打磨(完全并行)', indent=0.2)

# ================= 三、代码架构 =================
add_heading('三、代码架构总览', 1)
add_body('采用「管理器 Manager + 场景 Scene + 可交互物件 Interactable + 谜题 Puzzle + 全局状态 GameState」五件套。')
add_heading('系统层(A 负责)', 2)
for line in ['GameManager —— 全局入口、章节/循环调度',
             'SceneLoader —— 场景/房间加载与切换(含转场)',
             'SaveSystem —— 存档读写(JSON)',
             'GameState —— 全局标志位/物品/进度(可存档)',
             'InventorySystem —— 背包与物品使用(选中→点目标)',
             'PlayerController —— 主角老人行走(点击寻路 / 单屏可走带)',
             'CloseupView —— 特写视图(抽屉/信件细看,不换场景)',
             'AudioManager / UIManager / InputManager',
             '基类:InteractableBase(可点击物件,点击→走到交互点→触发)、PuzzleBase(谜题)、WalkableArea(房间可走区)、DialogueSystem(对白)']:
    add_body(line, bullet=True)
add_heading('内容层(B 负责)', 2)
for line in ['Room_00_Entrance / Room_01_xxx …… 各房间场景',
             '各种 Interactable 派生(门、抽屉、符号、道具)',
             '各章谜题 Puzzle 派生(拼图、序列、密码…)',
             '剧情触发器、过场 Timeline']:
    add_body(line, bullet=True)

add_heading('关键接口约定(A 先交付,B 依赖)——双人不冲突的命脉', 2)
code = doc.add_paragraph()
code.paragraph_format.left_indent = Inches(0.2)
run = code.add_run(
'GameState.GetFlag / SetFlag(key)      // 全局开关,进度都存这里\n'
'GameState.AddItem / HasItem(itemId)   // 物品\n'
'InteractableBase.OnClick()            // B 的可点击物都继承\n'
'PuzzleBase.OnSolved (UnityEvent)      // 谜题解开统一广播\n'
'SceneLoader.GoToRoom("Room_01")       // 场景切换,B 只调用'
)
set_font(run, 9.5, name='Consolas')
add_body('铁律:内容层只调用系统层暴露的公开接口,不直接读写内部字段。接口若要改,A 先在《接口契约》记一笔并通知 B。')

# ================= 三·五、主角操控系统 =================
add_heading('三·五、主角操控系统(点击寻路行走)', 1)
add_body('本作虽是绣湖类,但主角老人在场景内可行走(区别于经典绣湖的纯视点)。这是 2026-07-12 敲定的核心设计,详细接口见《接口契约》§7~§10。')
add_heading('玩家操控 = 四层', 2)
add_body('① 点击交互:光标 hover 高亮 → 点击触发。', bullet=True)
add_body('② 老人行走(玩家的"移动"):点空地→老人走过去;点物件→老人走到该物 interactPoint→到位后才触发 OnClick()。', bullet=True)
add_body('③ 视点切换:点门→SceneLoader.GoToRoom() 换房间;点抽屉/信件→CloseupView 弹特写层(不换场景)。', bullet=True)
add_body('④ 背包:选中物品→光标持物态→点目标物(InventorySystem.SelectedItem 判断)。', bullet=True)
add_heading('寻路方案(已定):单屏可走带', 2)
add_body('每房间一个 WalkableArea 圈出地面,老人直线移动并被 clamp 在可走区内,不绕障碍。零第三方依赖,够绣湖类单屏房间用;个别房间需绕障碍再单独升级为路点图。')
add_heading('各方新增职责', 2)
add_body('系统层 A:新增 PlayerController(行走/朝向/动画/深度排序)、WalkableArea(可走区)、CloseupView(特写);InteractableBase 改为"点击→走到交互点→触发"流程,新增 interactPoint / walkToBeforeInteract 字段。', bullet=True)
add_body('内容层 B:每房间摆一个 WalkableArea 圈地面;每个可交互物摆一个 interactPoint 空物体标"老人站哪儿点它"。', bullet=True)
add_body('美术:老人行走动画(至少 idle + 左右 walk 两套帧或骨骼动画),这是纯绣湖不需要的,需同步。', bullet=True)
code2 = doc.add_paragraph()
code2.paragraph_format.left_indent = Inches(0.2)
run = code2.add_run(
'PlayerController.WalkTo(pos / target, onArrive)   // 老人走到某处,到位回调\n'
'WalkableArea.ClampToArea(pos)                     // 把点击点拉回可走区\n'
'InventorySystem.Select / SelectedItem             // 背包选中→点目标\n'
'CloseupView.Open(id) / Close()                    // 特写细看,不换场景'
)
set_font(run, 9.5, name='Consolas')

# ================= 四、目录结构 =================
add_heading('四、Unity 工程目录结构(Assets 下)', 1)
add_body('规范的目录 = 双人不打架的物理基础。每个房间一个独立场景文件,是绣湖类避免场景冲突的关键。')
tree = doc.add_paragraph()
tree.paragraph_format.left_indent = Inches(0.15)
run = tree.add_run(
'Assets/_Project/\n'
'  ├─ Scenes/         Boot、MainMenu、Prologue/(每间房一个 .unity)\n'
'  ├─ Scripts/\n'
'  │    ├─ Core/      系统层(A):Managers / Save / Interaction / Dialogue\n'
'  │    ├─ Rooms/     内容层(B):按章分文件夹\n'
'  │    └─ Puzzles/   具体谜题脚本\n'
'  ├─ Art/            美术交付:Backgrounds / Sprites / UI / Animations\n'
'  ├─ Audio/          音乐、音效\n'
'  ├─ Prefabs/        预制体(可交互物做成 Prefab 复用)\n'
'  ├─ Data/           ScriptableObject:物品表/对白表/谜题配置/Keys.cs\n'
'  └─ Fonts/'
)
set_font(run, 9.5, name='Consolas')
add_heading('命名规范(强制)', 2)
add_body('场景:章前缀_序号_英文名 → P_01_Hall(P=序章 O=老年 M=中年 Y=青年)', bullet=True)
add_body('脚本类名 PascalCase;谜题以 Puzzle_ 结尾,交互物以 Interact_ 结尾', bullet=True)
add_body('Flag/物品 key 用 snake_case,集中放 Data/Keys.cs,禁止散落魔法字符串', bullet=True)
add_body('美术资源:bg_房间名、spr_物件名、ui_控件名', bullet=True)

# ================= 五、版本控制 =================
add_heading('五、版本控制与"防场景冲突"规范(重中之重)', 1)
add_body('Unity + Git 最大的坑:两人改同一个 .unity 或 .prefab → 合并几乎无法解决。以下规矩必须执行。')
add_heading('一次性设置', 2)
add_body('Git LFS:美术/音频/大二进制走 LFS(见 .gitattributes)', bullet=True)
add_body('Unity 强制文本序列化 + 可见 meta:Editor 设置里 Asset Serialization→Force Text,Version Control→Visible Meta Files', bullet=True)
add_body('配置 UnityYAMLMerge 智能合并场景/预制体', bullet=True)
add_heading('铁律(写进群公告)', 2)
add_body('一个场景/预制体同一时间只有一个人改,开工前在看板"认领"', bullet=True)
add_body('每个房间独立场景,天然把冲突面缩到最小', bullet=True)
add_body('系统层公共文件只有 A 改;B 需要新全局功能 → 提给 A', bullet=True)
add_body('每天开工先 git pull,收工前 git push,别攒一周', bullet=True)
add_body('提交信息带前缀:[core] / [room] / [puzzle] / [art] / [fix]', bullet=True)
add_heading('分支策略(轻量)', 2)
add_body('main(稳定可运行) ← dev(日常集成) ← feat/*(每人每功能一条分支)', indent=0.2)
add_body('完成→合进 dev→大家 pull;里程碑稳定后 dev 合 main 打 tag(如 v0.1-prologue)。小团队可简化为两人各推 dev,但必须勤 pull。')

# ================= 六、准备工作 =================
add_heading('六、准备工作清单(开工前 1~2 天完成)', 1)
add_heading('程序员 A(系统程序)', 2)
for x in ['装 Unity Hub + Unity 2022.3 LTS(含 Windows Build、2D 模块)',
          '新建 Unity 工程(2D Core 模板),工程名 LostGoddess',
          '设置 Force Text + Visible Meta',
          '建立 Assets/_Project 目录结构',
          'git init + 放入 .gitignore / .gitattributes',
          'git lfs install,推到 Gitee/GitHub',
          '导入 DOTween(补间动画)、可选 Newtonsoft Json',
          '搭 Boot/MainMenu 场景骨架 + 空 Manager + 接口基类空实现',
          '首次推送,让 B clone 验证能跑']:
    add_body('☐ ' + x, indent=0.15)
add_heading('程序员 B(内容程序)', 2)
for x in ['装好同版本 Unity 2022.3 LTS(版本号必须与 A 完全一致)',
          'Clone 仓库,确认能打开、能进 Boot 场景运行',
          '熟悉接口基类,做"点击物件→播音效→设 flag→开门"最小交互原型',
          '与美术对接占位图规格(分辨率、命名、锚点)']:
    add_body('☐ ' + x, indent=0.15)
add_heading('需要向其他三位同学索取的输入', 2)
add_body('策划:序章流程图、房间数量、每个谜题解法与前后置关系、物品清单', bullet=True)
add_body('美术:占位图或首批背景/物件 sprite、分辨率与命名规范确认', bullet=True)
add_body('文案:序章旁白/对白文本(草稿也行,先让对白系统跑起来)', bullet=True)

# ================= 七、序章排期 =================
add_heading('七、序章 10 天开发排期(示例,按实际调整)', 1)
add_table(
    ['天', '程序员 A(系统)', '程序员 B(内容)', '里程碑'],
    [['D1', '建工程/Git/LFS/目录,推首版', 'Clone 跑通,读接口文档', '仓库可协作'],
     ['D2', 'GameManager + SceneLoader(含淡入淡出)', '最小交互原型(点击→反馈)', '能切场景'],
     ['D3', 'PlayerController + WalkableArea(点击→老人走过去)', '搭第一个房间 + 占位图 + WalkableArea', '主角能走'],
     ['D4', 'GameState + SaveSystem(存/读 flag 与物品)', '摆 interactPoint,门/抽屉(走近→触发)', '存档+走近交互通'],
     ['D5', 'InteractableBase/PuzzleBase 定稿 + CloseupView', '第一个完整谜题(特写里,密码/序列)', '谜题闭环'],
     ['D6', 'InventorySystem(选中→点目标)+ 物品使用', '接入文案,房间 2', '背包可用'],
     ['D7', 'DialogueSystem + UIManager + AudioManager', '房间 3 + 谜题 2', '剧情/音画可播'],
     ['D8', '主菜单/存档槽/章节循环钩子', '剩余房间铺量', '序章骨架全'],
     ['D9', '联调、修系统 bug、性能', '联调、补交互、过场', '可通关'],
     ['D10', '打包出 Demo + 复盘', '打磨手感/音画同步', '序章 Demo']]
)
add_body('时间循环(同一场景在青年/中年/老年三态切换)是本作特色,序章先埋接口不做全:GameManager 预留 CurrentEra 枚举与场景变体加载钩子,等老年章启用,后三章即可复用序章框架。')

# ================= 八、复用建议 =================
add_heading('八、给后续章节的复用建议', 1)
add_body('序章做出的系统层 = 全项目地基,后三章基本只写内容层,产能显著提升。', bullet=True)
add_body('"同一房间三个时代版本"用同场景 + 按 CurrentEra 切换 sprite/物件启用实现,避免建三份场景。', bullet=True)
add_body('每章沿用相同目录/命名/分支规范,新程序员也能快速上手。', bullet=True)

# ================= 附录 =================
add_heading('附录:随项目已备文件', 1)
add_body('项目骨架已建在 C:\\Users\\yuan\\lost-goddess-game\\,clone 即可用:')
add_body('.gitignore(Unity 专用) / .gitattributes(LFS + 场景合并)', bullet=True)
add_body('README.md(团队须知与环境版本)', bullet=True)
add_body('docs/接口契约.md(A 维护、B 依赖的接口清单,改动即更新)', bullet=True)
add_body('docs/程序开发计划方案.md(本文档的完整 Markdown 版)', bullet=True)

out = 'C:/Users/yuan/Desktop/失落的女神_程序开发计划方案_20260712.docx'
doc.save(out)
print('已生成:', out)
